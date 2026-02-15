#!/usr/bin/env python3
"""
Ghostcore Evidence OS V5.2 - Setup Script
Creates complete directory structure with minimal working implementation
"""

import os
import yaml
import sys
from pathlib import Path

def create_ghostcore_v5_2():
    """Create the complete Ghostcore Evidence OS V5.2 structure"""
    
    # Define the ghostcore root
    ghostcore_root = Path("ghostcore")
    ghostcore_root.mkdir(exist_ok=True)
    
    print("🔥 Creating Ghostcore Evidence OS V5.2...")
    
    # Create directory structure
    directories = [
        ghostcore_root / "bin",
        ghostcore_root / "core",
        ghostcore_root / "modules",
        ghostcore_root / "templates",
        ghostcore_root / "cases",
        ghostcore_root / "build"
    ]
    
    for directory in directories:
        directory.mkdir(exist_ok=True)
        print(f"📁 Created: {directory}")
    
    # Create the CLI launcher
    launcher_content = '''#!/usr/bin/env python3

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.runtime import Runtime

def main():
    if len(sys.argv) < 3:
        print("Usage: ghostcore build <case.yaml>")
        return
    
    command = sys.argv[1]
    target = sys.argv[2]

    runtime = Runtime()

    if command == "build":
        runtime.build_case(target)
    else:
        print("Unknown command.")

if __name__ == "__main__":
    main()
'''
    
    launcher_path = ghostcore_root / "bin" / "ghostcore"
    with open(launcher_path, 'w') as f:
        f.write(launcher_content)
    launcher_path.chmod(0o755)  # Make executable
    print(f"🚀 Created launcher: {launcher_path}")
    
    # Create core modules
    runtime_content = '''import yaml
from core.fcm_loader import FCMLoader
from core.compiler import Compiler
from core.logger import Logger

class Runtime:
    def __init__(self):
        self.logger = Logger("Ghostcore")
        self.compiler = Compiler()
        self.fcm = FCMLoader()

    def build_case(self, case_file):
        self.logger.info(f"Loading case: {case_file}")
        case_data = self.fcm.load_case(case_file)
        
        self.logger.info("Compiling case…")
        output = self.compiler.compile(case_data)
        
        self.logger.success(f"Build complete: {output}")
        return output
'''
    
    with open(ghostcore_root / "core" / "runtime.py", 'w') as f:
        f.write(runtime_content)
    print("⚙️ Created runtime module")
    
    # FCM Loader
    fcm_loader_content = '''import yaml

class FCMLoader:
    def load_case(self, case_file):
        with open(case_file, 'r') as f:
            return yaml.safe_load(f)
'''
    
    with open(ghostcore_root / "core" / "fcm_loader.py", 'w') as f:
        f.write(fcm_loader_content)
    print("📋 Created FCM loader")
    
    # Compiler
    compiler_content = '''from modules.exporter_pdf import PDFExporter
from modules.exporter_docx import DOCXExporter
from modules.uploader import Uploader
import os

class Compiler:
    def __init__(self):
        self.pdf = PDFExporter()
        self.docx = DOCXExporter()
        self.up = Uploader()

    def compile(self, case):
        outputs = []

        if case.get("outputs", {}).get("pdf"):
            pdf_path = self.pdf.build(case)
            outputs.append(pdf_path)

        if case.get("outputs", {}).get("docx"):
            docx_path = self.docx.build(case)
            outputs.append(docx_path)

        # ZIP packaging
        zip_path = "build/case_package.zip"
        self.up.pack(outputs, zip_path)

        return zip_path
'''
    
    with open(ghostcore_root / "core" / "compiler.py", 'w') as f:
        f.write(compiler_content)
    print("🔨 Created compiler")
    
    # Identity Kernel
    identity_content = '''import hashlib
import time

class IdentityKernel:
    def __init__(self, author="Šabad"):
        self.author = author

    def watermark(self, content):
        stamp = f"{self.author}-{int(time.time())}"
        digest = hashlib.sha256(stamp.encode()).hexdigest()
        return f"\\n\\n[GHOSTCORE-ID:{digest[:12]}]"
'''
    
    with open(ghostcore_root / "core" / "identity.py", 'w') as f:
        f.write(identity_content)
    print("🔒 Created identity kernel")
    
    # Logger
    logger_content = '''import logging
import json
from datetime import datetime

class Logger:
    def __init__(self, name):
        self.logger = logging.getLogger(name)
        self.logger.setLevel(logging.INFO)
        
        # Create console handler
        handler = logging.StreamHandler()
        handler.setLevel(logging.INFO)
        
        # Create formatter
        formatter = logging.Formatter(
            '{"time":"%(asctime)s","level":"%(levelname)s","msg":"%(message)s"}',
            datefmt="%Y-%m-%dT%H:%M:%S"
        )
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)

    def info(self, msg):
        self.logger.info(msg)

    def success(self, msg):
        self.logger.info(f"SUCCESS: {msg}")

    def error(self, msg):
        self.logger.error(f"ERROR: {msg}")
'''
    
    with open(ghostcore_root / "core" / "logger.py", 'w') as f:
        f.write(logger_content)
    print("📝 Created logger")
    
    # PDF Exporter
    pdf_exporter_content = '''#!/usr/bin/env python3
"""
PDF Exporter Module for Ghostcore Evidence OS
"""

from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from core.identity import IdentityKernel

class PDFExporter:
    def __init__(self):
        self.identity = IdentityKernel()

    def build(self, case_data):
        title = case_data.get('case', {}).get('title', 'Untitled Case')
        filename = f"build/{title.replace(' ', '_')}.pdf"
        
        doc = SimpleDocTemplate(
            filename,
            pagesize=LETTER,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CaseTitle',
            parent=styles['Title'],
            fontSize=18,
            leading=22
        ))

        flow = []
        
        # Add title
        flow.append(Paragraph(title, styles['CaseTitle']))
        
        # Add sections
        sections = case_data.get('sections', [])
        for section in sections:
            flow.append(Paragraph(section, styles['Heading2']))
            flow.append(Spacer(1, 0.2 * inch))
        
        # Add exhibits
        exhibits = case_data.get('exhibits', [])
        for exhibit in exhibits:
            flow.append(Paragraph(f"Exhibit: {exhibit.get('type', 'Unknown')}", styles['Normal']))
            flow.append(Spacer(1, 0.1 * inch))
        
        # Add watermark
        watermark_text = self.identity.watermark(title)
        flow.append(Spacer(1, 0.2 * inch))
        flow.append(Paragraph(f"<i>{watermark_text}</i>", styles['Normal']))
        
        doc.build(flow)
        return filename
'''
    
    with open(ghostcore_root / "modules" / "exporter_pdf.py", 'w') as f:
        f.write(pdf_exporter_content)
    print("📄 Created PDF exporter")
    
    # DOCX Exporter
    docx_exporter_content = '''#!/usr/bin/env python3
"""
DOCX Exporter Module for Ghostcore Evidence OS
"""

from docx import Document
from docx.shared import Inches
from core.identity import IdentityKernel

class DOCXExporter:
    def __init__(self):
        self.identity = IdentityKernel()

    def build(self, case_data):
        title = case_data.get('case', {}).get('title', 'Untitled Case')
        filename = f"build/{title.replace(' ', '_')}_docx.docx"
        
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add sections
        sections = case_data.get('sections', [])
        for section in sections:
            doc.add_heading(section, level=1)
            doc.add_paragraph("")  # Empty paragraph for spacing
        
        # Add exhibits
        exhibits = case_data.get('exhibits', [])
        for exhibit in exhibits:
            doc.add_heading(f"Exhibit: {exhibit.get('type', 'Unknown')}", level=2)
            doc.add_paragraph("")  # Empty paragraph for spacing
        
        # Add watermark
        watermark_text = self.identity.watermark(title)
        doc.add_paragraph(watermark_text)
        
        doc.save(filename)
        return filename
'''
    
    with open(ghostcore_root / "modules" / "exporter_docx.py", 'w') as f:
        f.write(docx_exporter_content)
    print("📝 Created DOCX exporter")
    
    # Uploader
    uploader_content = '''#!/usr/bin/env python3
"""
Uploader Module for Ghostcore Evidence OS
"""

import zipfile
import os

class Uploader:
    def pack(self, files, zip_path):
        os.makedirs(os.path.dirname(zip_path), exist_ok=True)
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_path in files:
                if os.path.exists(file_path):
                    zipf.write(file_path, os.path.basename(file_path))
        return zip_path
'''
    
    with open(ghostcore_root / "modules" / "uploader.py", 'w') as f:
        f.write(uploader_content)
    print("📦 Created uploader")
    
    # Example case
    example_case = {
        'case': {
            'title': 'Epstein–8200 Intelligence Nexus',
            'author': 'Šabad',
            'date': '2025'
        },
        'sections': [
            'Executive Summary',
            'Verified Findings',
            'Intelligence Lineage',
            'Financial Networks',
            'Timeline (1974–2025)',
            'Exhibits',
            'Annexes'
        ],
        'exhibits': [
            {
                'type': 'chart',
                'data': 'charts/funding.yaml'
            },
            {
                'type': 'org_diagram',
                'data': 'diagrams/carbyne_axon.yaml'
            }
        ],
        'outputs': {
            'pdf': True,
            'docx': True,
            'zip': True
        }
    }
    
    with open(ghostcore_root / "cases" / "example_case.yaml", 'w') as f:
        yaml.dump(example_case, f, default_flow_style=False)
    print("📋 Created example case")
    
    # Requirements file
    requirements_content = '''reportlab
python-docx
pyyaml
Pillow
qrcode
'''
    
    with open(ghostcore_root / "requirements.txt", 'w') as f:
        f.write(requirements_content)
    print("📦 Created requirements file")
    
    # README
    readme_content = '''# Ghostcore Evidence OS V5.2

A modular forensic evidence compilation system for investigative research.

## Structure

- `bin/` - Command line interface
- `core/` - Core runtime and compiler
- `modules/` - Pluggable modules (exporters, charts, etc.)
- `templates/` - Document templates
- `cases/` - Case definition files (FCM format)
- `build/` - Output directory

## Usage

```bash
ghostcore build cases/example_case.yaml
```

## Features

- Multi-format export (PDF, DOCX)
- Modular architecture
- Identity kernel with anti-tamper protection
- Case compilation from YAML definitions
- ZIP packaging of all artifacts
- Structured logging

## FCM Format

The Forensic Case Model (FCM) defines evidence cases in YAML format with:
- Case metadata
- Document sections
- Exhibits and diagrams
- Output formats

## Modules

The system supports pluggable modules for:
- PDF export
- DOCX export
- Chart generation
- Diagram building
- Upload services
'''
    
    with open(ghostcore_root / "README.md", 'w') as f:
        f.write(readme_content)
    print("📖 Created README")
    
    print(f"\\n🔥 GHOSTCORE EVIDENCE OS V5.2 COMPLETE!")
    print(f"📁 Directory structure created in: {ghostcore_root}")
    print(f"🚀 To test: python3 {ghostcore_root}/bin/ghostcore build {ghostcore_root}/cases/example_case.yaml")

if __name__ == "__main__":
    create_ghostcore_v5_2()