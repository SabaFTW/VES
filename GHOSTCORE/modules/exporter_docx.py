#!/usr/bin/env python3
"""
DOCX Exporter Module for Ghostcore Evidence OS
"""

from docx import Document
from docx.shared import Inches
from core.identity import IdentityKernel

class DOCXExporter:
    def __init__(self):
        self.identity = IdentityKernel()

    def build(self, case_data):
        title = case_data.get('case', {}).get('title', 'Untitled Case')
        filename = f"build/{title.replace(' ', '_')}_docx.docx"
        
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add sections
        sections = case_data.get('sections', [])
        for section in sections:
            doc.add_heading(section, level=1)
            doc.add_paragraph("")  # Empty paragraph for spacing
        
        # Add exhibits
        exhibits = case_data.get('exhibits', [])
        for exhibit in exhibits:
            doc.add_heading(f"Exhibit: {exhibit.get('type', 'Unknown')}", level=2)
            doc.add_paragraph("")  # Empty paragraph for spacing
        
        # Add watermark
        watermark_text = self.identity.watermark(title)
        doc.add_paragraph(watermark_text)
        
        doc.save(filename)
        return filename
